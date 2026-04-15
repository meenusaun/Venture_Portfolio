import streamlit as st
import pandas as pd
import json
import os
import io
import fitz  # PyMuPDF
from docx import Document as DocxDocument
from anthropic import Anthropic

# Google Drive imports
try:
    from googleapiclient.discovery import build
    from googleapiclient.http import MediaIoBaseDownload
    from google.oauth2 import service_account
    GDRIVE_AVAILABLE = True
except ImportError:
    GDRIVE_AVAILABLE = False

# ─────────────────────────────────────────────────────────────────────────────
# CONFIG
# ─────────────────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Venture Intelligence | NEN Resources Network",
    page_icon="🚀",
    layout="wide",
    initial_sidebar_state="expanded"
)

client = Anthropic()

for key, default in [
    ("chat_messages", []),
    ("analysis_cache", {}),
    ("extraction_cache", {}),
]:
    if key not in st.session_state:
        st.session_state[key] = default

# ─────────────────────────────────────────────────────────────────────────────
# STYLES
# ─────────────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
.badge-active  { background:#d1fae5; color:#065f46; border:1px solid #6ee7b7; padding:2px 10px; border-radius:20px; font-size:11px; font-weight:600; }
.badge-alumni  { background:#ede9fe; color:#5b21b6; border:1px solid #c4b5fd; padding:2px 10px; border-radius:20px; font-size:11px; font-weight:600; }
.badge-idea    { background:#fef3c7; color:#92400e; border:1px solid #fcd34d; padding:2px 10px; border-radius:20px; font-size:11px; font-weight:600; }
.badge-mvp     { background:#dbeafe; color:#1e40af; border:1px solid #93c5fd; padding:2px 10px; border-radius:20px; font-size:11px; font-weight:600; }
.badge-growth  { background:#d1fae5; color:#065f46; border:1px solid #6ee7b7; padding:2px 10px; border-radius:20px; font-size:11px; font-weight:600; }
.nen-card        { background:#f8fafc; border:1px solid #e2e8f0; border-radius:12px; padding:20px 24px; margin-bottom:16px; }
.nen-card-accent { background:#eff6ff; border:1px solid #bfdbfe; border-radius:12px; padding:20px 24px; margin-bottom:16px; }
.ai-insight { background:#eff6ff; border-left:3px solid #2563eb; padding:14px 18px; border-radius:0 10px 10px 0; font-size:14px; line-height:1.7; color:#1e3a5f; margin:10px 0; }
.journey-arrow { color:#2563eb; font-size:18px; margin:0 8px; }
</style>
""", unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# FILE CLASSIFICATION & READING
# ─────────────────────────────────────────────────────────────────────────────
FILE_TYPE_KEYWORDS = {
    "deck":     ["deck", "pitch", "presentation"],
    "vp":       ["vp1", "vp2", "vp3", "vp4", "vp5", "vp_", "_vp", "vp2 ",
                 "opportunity assessment", "opportunity_assessment",
                 "opportunity assessment transcript",
                 "deep dive", "deep_dive", "deepdive",
                 "vp call", "vp-call", "vp transcript", "vp1 transcript",
                 "vp2 transcript", "vp3 transcript"],
    "expert":   ["expert_connect", "expert connect", "expertconnect",
                 "expert_session", "expert session", "mentor"],
    "panelist": ["panel evaluation", "panel score", "panelist", "panellist",
                 "panel transcript", "discussion with", "panel "],
    "context":  ["growth plan", "impact dashboard", "score sheet",
                 "evaluation form", "strategic growth", "blueprint"],
}
SUPPORTED_EXTS = {".pdf", ".docx", ".doc", ".txt", ".md", ".xlsx", ".xls"}


def classify_file(filename: str) -> str:
    name = filename.lower()
    for ftype, keywords in FILE_TYPE_KEYWORDS.items():
        if any(kw in name for kw in keywords):
            return ftype
    return "other"


def read_file_text(path: str, max_chars: int = 12000) -> str:
    if not path or not os.path.exists(path):
        return ""
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            doc = fitz.open(path)
            return "\n".join(page.get_text() for page in doc)[:max_chars]
        elif ext in [".docx", ".doc"]:
            doc = DocxDocument(path)
            return "\n".join(p.text for p in doc.paragraphs)[:max_chars]
        elif ext in [".txt", ".md"]:
            return open(path, encoding="utf-8").read()[:max_chars]
    except Exception as e:
        return f"[Could not read: {e}]"
    return ""


def read_file_bytes(file_bytes: bytes, filename: str, max_chars: int = 12000) -> str:
    ext = os.path.splitext(filename)[1].lower()
    try:
        if ext == ".pdf":
            doc = fitz.open(stream=file_bytes, filetype="pdf")
            return "\n".join(page.get_text() for page in doc)[:max_chars]
        elif ext in [".docx", ".doc"]:
            doc = DocxDocument(io.BytesIO(file_bytes))
            return "\n".join(p.text for p in doc.paragraphs)[:max_chars]
        elif ext in [".txt", ".md"]:
            return file_bytes.decode("utf-8", errors="ignore")[:max_chars]
        elif ext in [".xlsx", ".xls"]:
            try:
                dfs = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
                parts = []
                for sheet, df in dfs.items():
                    parts.append(f"[Sheet: {sheet}]\n{df.to_string(index=False)}")
                return "\n\n".join(parts)[:max_chars]
            except Exception:
                return "[Excel file — could not parse]"
    except Exception as e:
        return f"[Could not read: {e}]"
    return ""


# ─────────────────────────────────────────────────────────────────────────────
# GOOGLE DRIVE
# ─────────────────────────────────────────────────────────────────────────────
@st.cache_resource
def get_gdrive_service():
    if not GDRIVE_AVAILABLE:
        return None
    try:
        creds_dict = dict(st.secrets["gdrive_service_account"])
        if "private_key" in creds_dict:
            creds_dict["private_key"] = creds_dict["private_key"].replace("\\n", "\n")
        creds = service_account.Credentials.from_service_account_info(
            creds_dict, scopes=["https://www.googleapis.com/auth/drive.readonly"]
        )
        return build("drive", "v3", credentials=creds, cache_discovery=False)
    except Exception:
        return None


def gdrive_find_venture_folder(service, venture_name: str, parent_name: str = "Venture_Docs") -> tuple:
    try:
        parent_id = None
        for pname in [parent_name, parent_name.replace("_", " ")]:
            res = service.files().list(
                q=f"name=\'{pname}\' and mimeType=\'application/vnd.google-apps.folder\' and trashed=false",
                fields="files(id, name)", pageSize=5
            ).execute()
            if res.get("files"):
                parent_id = res["files"][0]["id"]
                break
        if not parent_id:
            return None, f"Parent \'{parent_name}\' not found. Share Venture_Docs with the service account."
        res_all = service.files().list(
            q=f"mimeType=\'application/vnd.google-apps.folder\' and \'{parent_id}\' in parents and trashed=false",
            fields="files(id, name)", pageSize=50
        ).execute()
        folders = res_all.get("files", [])
        names = [f["name"] for f in folders]
        for f in folders:
            if f["name"].lower().strip() == venture_name.lower().strip():
                return f["id"], "ok"
        return None, f"\'{venture_name}\' not found. Available: {', '.join(names) or 'none'}"
    except Exception as e:
        return None, f"Drive error: {e}"


def gdrive_list_files(service, folder_id: str) -> list:
    try:
        res = service.files().list(
            q=f"\'{folder_id}\' in parents and trashed=false",
            fields="files(id, name, mimeType)", pageSize=50
        ).execute()
        return [f for f in res.get("files", [])
                if os.path.splitext(f["name"])[1].lower() in SUPPORTED_EXTS]
    except Exception:
        return []


def gdrive_download(service, file_id: str) -> bytes:
    try:
        buf = io.BytesIO()
        downloader = MediaIoBaseDownload(buf, service.files().get_media(fileId=file_id))
        done = False
        while not done:
            _, done = downloader.next_chunk()
        return buf.getvalue()
    except Exception:
        return b""


def load_venture_files(venture_name: str, folder_path: str = "") -> dict:
    """Load all files for a venture from Google Drive (primary) or local path (fallback)."""
    result = {
        "files": [], "deck": [], "vp": [], "expert": [],
        "panelist": [], "other": [],
        "source": "none", "folder_exists": False, "total": 0, "debug": ""
    }
    service = get_gdrive_service()
    lookup = folder_path.strip() if folder_path.strip() else venture_name.strip()

    if service and lookup:
        folder_id, debug = gdrive_find_venture_folder(service, lookup)
        result["debug"] = debug
        if folder_id:
            result["folder_exists"] = True
            result["source"] = "gdrive"
            drive_files = gdrive_list_files(service, folder_id)

            # Deduplicate: if exact same base name exists as both .docx AND .pdf,
            # keep only the PDF. Different filenames are always kept even if similar.
            seen_bases = {}
            for f in drive_files:
                base = os.path.splitext(f["name"])[0].lower().strip()
                ext  = os.path.splitext(f["name"])[1].lower()
                ftype = classify_file(f["name"])
                key  = (base, ftype)  # only dedup if same name AND same type
                if key not in seen_bases:
                    seen_bases[key] = f
                else:
                    existing_ext = os.path.splitext(seen_bases[key]["name"])[1].lower()
                    if ext == ".pdf" and existing_ext in [".docx", ".doc"]:
                        seen_bases[key] = f
            deduped_files = list(seen_bases.values())

            for f in deduped_files:
                fname = f["name"]
                ftype = classify_file(fname)
                raw   = gdrive_download(service, f["id"])
                text  = read_file_bytes(raw, fname, 10000 if ftype in ("deck","context") else 14000) if raw else ""
                entry = {"name": fname, "type": ftype, "text": text}
                result.setdefault(ftype, []).append(entry)
                result["files"].append(entry)
            result["total"] = len(result["files"])
            return result

    if folder_path and os.path.isdir(folder_path):
        result["folder_exists"] = True
        result["source"] = "local"
        for fname in sorted(os.listdir(folder_path)):
            if os.path.splitext(fname)[1].lower() not in SUPPORTED_EXTS:
                continue
            ftype = classify_file(fname)
            text = read_file_text(os.path.join(folder_path, fname),
                                  10000 if ftype == "deck" else 14000)
            entry = {"name": fname, "type": ftype, "text": text}
            result[ftype].append(entry)
            result["files"].append(entry)
        result["total"] = len(result["files"])
    return result


# ─────────────────────────────────────────────────────────────────────────────
# EXCEL (slim — basics only)
# ─────────────────────────────────────────────────────────────────────────────
@st.cache_data(show_spinner=False)
def load_excel(file_bytes: bytes) -> pd.DataFrame:
    xl = pd.read_excel(io.BytesIO(file_bytes), sheet_name=None)
    for name in ["Ventures", "ventures"] + list(xl.keys()):
        if name in xl:
            df = xl[name].copy()
            df.columns = [str(c).strip() for c in df.columns]
            return df
    return pd.DataFrame()


# ─────────────────────────────────────────────────────────────────────────────
# AI — STEP 1: EXTRACT SESSIONS FROM FILES
# ─────────────────────────────────────────────────────────────────────────────
def ai_extract_sessions(venture_name: str, loaded_files: dict) -> dict:
    # Separate files by category — process context/deck separately from sessions
    session_files  = [e for e in loaded_files["files"] if e["type"] not in ("deck","context")]
    context_files  = [e for e in loaded_files["files"] if e["type"] in ("deck","context","other")]

    if not session_files and not context_files:
        return {"sessions": [], "extraction_notes": "No files found in folder."}

    # Build session text — cap each file at 3500 chars to fit more files
    # Map internal type to display label used in JSON output
    TYPE_LABEL = {
        "vp":       "VP Session",
        "expert":   "Expert Session",
        "panelist": "Panelist Call",
        "other":    "Other",
    }
    files_text = ""
    for entry in session_files:
        forced_type = TYPE_LABEL.get(entry["type"], "Other")
        files_text += f"\n\n{'='*60}\nFILE: {entry['name']}\nSESSION TYPE (MANDATORY — use exactly this): {forced_type}\n{'='*60}\n"
        files_text += (entry["text"] or "[empty]")[:3500]

    # Build context text from deck/other docs
    context_text = ""
    for entry in context_files:
        context_text += f"\n\n--- CONTEXT FILE: {entry['name']} ---\n"
        context_text += (entry["text"] or "[empty]")[:2000]

    if not files_text.strip():
        return {"sessions": [], "extraction_notes": "Only deck/context files found — no session transcripts."}

    prompt = f"""You are analyzing all documents for startup venture "{venture_name}".
You will extract two things:
1. Venture context (brief, problem statement, solution) from context/deck documents
2. Every individual meeting or session from transcript files

FILE CLASSIFICATION (already applied to filenames — use the (type:) label shown):
- type: expert  → Expert Session transcript
- type: vp      → VP Session transcript
- type: panelist → Panelist Call transcript  
- type: deck    → Pitch deck (venture info only, not a session)
- type: context → Supporting document (growth plan, scorecard etc.)
- type: other   → Read and extract whatever sessions you find

IMPORTANT: Even if a file is type: context or other, still extract any session content if present.

CONTEXT DOCUMENTS (for venture info):
{context_text[:3000]}

SESSION TRANSCRIPT FILES:
{files_text[:18000]}

RULES:
- Do NOT invent sessions — only extract what is clearly present
- Each file may contain one or multiple sessions
- For deck/brief files: extract venture_brief and problem_statement (not as a session)
- For session files: extract as individual sessions
- CRITICAL: The "type" field in your JSON MUST exactly match the "SESSION TYPE (MANDATORY)" label shown above each file. Do not override it based on content — trust the filename classification.
- Valid type values: "VP Session", "Expert Session", "Panelist Call", "Other"

Return ONLY valid JSON, no markdown:
{{
  "extraction_notes": "brief summary: how many sessions found, what types, any issues",
  "venture_brief": "2-3 sentence description of what this venture does, extracted from files (or Not found)",
  "problem_statement": "the core problem this venture is solving, extracted from files (or Not found)",
  "solution_summary": "what the venture's solution is, extracted from files (or Not found)",
  "founder_background": "any founder/team info found in files (or Not found)",
  "stage_inferred": "current stage inferred from sessions and documents: Idea / MVP / Growth / Scaling (pick the best fit based on what you read)",
  "stage_rationale": "one sentence explaining why you assessed this stage",
  "cohort": "cohort or program batch name if mentioned anywhere in files, else Not found",
  "current_status": "Active / Alumni / Stalled — inferred from recency of sessions and progress signals",
  "sessions": [
    {{
      "session_id": "S01",
      "type": "VP Session / Expert Session / Panelist Call / Other",
      "date": "date or Unknown",
      "participants": "names and roles mentioned",
      "problems_discussed": ["problem 1", "problem 2"],
      "action_items": ["action 1", "action 2"],
      "notes": "2-3 sentence summary of what happened in this session",
      "source_file": "filename",
      "effectiveness_rating": 1-5
    }}
  ]
}}"""

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=4000,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = response.content[0].text.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw.strip().rstrip("```").strip())


# ─────────────────────────────────────────────────────────────────────────────
# AI — STEP 2: FULL ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
def ai_analyze_venture(venture_row: pd.Series, loaded_files: dict, extracted: dict) -> dict:
    venture_info = venture_row.to_dict()

    # Prefer extracted venture info from files; fall back to Excel
    venture_brief    = extracted.get("venture_brief","") or venture_info.get("Description","")
    problem_stmt     = extracted.get("problem_statement","") or "Not extracted"
    solution_summary = extracted.get("solution_summary","") or "Not extracted"
    founder_bg       = extracted.get("founder_background","") or venture_info.get("Founder_Name","")

    # Deck text as supplementary context
    deck_parts = [f["text"] for f in loaded_files.get("deck", []) if f.get("text")]
    deck_text = "\n\n".join(deck_parts)
    deck_section = f"\nPITCH DECK (additional context):\n{deck_text[:4000]}\n" if deck_text else ""
    has_deck = "true" if deck_text else "false"

    sessions = extracted.get("sessions", [])
    total = len(sessions)
    sids = ", ".join(s.get("session_id", "") for s in sessions)

    stage_inferred  = extracted.get("stage_inferred","")
    stage_rationale = extracted.get("stage_rationale","")
    current_status  = extracted.get("current_status","Active")
    cohort          = extracted.get("cohort","") or venture_info.get("Cohort","")

    prompt = f"""Startup accelerator analyst. Produce a comprehensive intelligence report for this venture.

VENTURE (from files):
Name: {venture_info.get('Venture_Name','')}
Founder: {founder_bg}
Sector: {venture_info.get('Sector','')}
Current Stage (inferred from files): {stage_inferred} — {stage_rationale}
Cohort: {cohort}
Status: {current_status}
Venture Brief: {venture_brief}
Problem Statement: {problem_stmt}
Solution: {solution_summary}
{deck_section}
SESSIONS ({total}, IDs: {sids}):
{json.dumps(sessions, indent=2)[:9000]}

STRICT RULES:
- meeting_effectiveness: EXACTLY {total} entries, one per session_id ({sids})
- Use exact session_ids as meeting_ref
- health_before = state BEFORE any sessions (from starting stage + deck)
- health_after = state NOW after all sessions

Return ONLY valid JSON:
{{
  "deck_insights": {{
    "has_deck": {has_deck},
    "problem_statement": "from deck or Not available",
    "solution_summary": "from deck or Not available",
    "deck_strengths": ["s1"],
    "deck_gaps": ["g1"],
    "deck_vs_reality": "one sentence"
  }},
  "key_problems": [
    {{"problem": "...", "source": "session id or deck", "status": "resolved/ongoing/unaddressed"}}
  ],
  "journey_narrative": "2-3 paragraph story of the venture's journey through the program",
  "stage_progression": {{
    "from": "earliest stage signals seen in first sessions",
    "to": "current stage based on latest sessions",
    "momentum": "strong/moderate/slow/stalled",
    "momentum_reason": "one sentence"
  }},
  "meeting_effectiveness": [
    {{
      "meeting_ref": "exact session_id",
      "type": "VP Session/Expert Session/Panelist Call",
      "date": "date",
      "effectiveness_score": 1-5,
      "impact": "what changed",
      "is_most_impactful": true/false
    }}
  ],
  "most_impactful_meeting": {{"ref": "id", "type": "type", "reason": "why"}},
  "biggest_unresolved_problem": "one sentence",
  "health_before": {{"score": 1-10, "rationale": "one sentence"}},
  "health_after":  {{"score": 1-10, "rationale": "one sentence"}},
  "next_priority": "top recommended action"
}}"""

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=3000,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = response.content[0].text.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw.strip().rstrip("```").strip())


# ─────────────────────────────────────────────────────────────────────────────
# AI — PORTFOLIO
# ─────────────────────────────────────────────────────────────────────────────
def ai_portfolio_analysis(ventures_df: pd.DataFrame) -> dict:
    rows = []
    for _, row in ventures_df.iterrows():
        vid = str(row.get("Venture_ID", ""))
        cached = st.session_state.analysis_cache.get(vid, {})
        ext_cache = st.session_state.extraction_cache.get(vid, {})
        rows.append({
            "name": row.get("Venture_Name", ""),
            "sector": row.get("Sector", ""),
            "stage_inferred": ext_cache.get("stage_inferred", "Unknown"),
            "status": ext_cache.get("current_status", "Active"),
            "health_before": cached.get("health_before", {}).get("score", "?"),
            "health_after":  cached.get("health_after",  {}).get("score", "?"),
            "sessions": len(ext_cache.get("sessions", [])),
        })

    prompt = f"""NEN accelerator portfolio analyst. Analyze this cohort.

{json.dumps(rows, indent=2)}

Return ONLY valid JSON:
{{
  "portfolio_health": {{"score": 1-10, "summary": "2-3 sentences"}},
  "cohort_narrative": "3-4 sentence story",
  "needs_attention": [{{"name": "...", "reason": "..."}}],
  "best_performing": [{{"name": "...", "reason": "..."}}],
  "meeting_type_effectiveness": {{
    "most_effective_type": "...", "reason": "...",
    "least_effective_type": "...", "least_effective_reason": "..."
  }},
  "portfolio_recommendations": ["rec1", "rec2", "rec3"],
  "stage_distribution": {{"summary": "one sentence"}}
}}"""

    response = client.messages.create(
        model="claude-sonnet-4-20250514",
        max_tokens=1500,
        messages=[{"role": "user", "content": prompt}]
    )
    raw = response.content[0].text.strip()
    if raw.startswith("```"):
        raw = raw.split("```")[1]
        if raw.startswith("json"):
            raw = raw[4:]
    return json.loads(raw.strip().rstrip("```").strip())


# ─────────────────────────────────────────────────────────────────────────────
# UI HELPERS
# ─────────────────────────────────────────────────────────────────────────────
def health_color(score):
    try:
        s = int(score)
        if s >= 8: return "#16a34a"
        if s >= 6: return "#2563eb"
        if s >= 4: return "#d97706"
        return "#dc2626"
    except Exception:
        return "#94a3b8"


def stage_badge(stage: str) -> str:
    s = str(stage).lower()
    if "idea" in s:   return '<span class="badge-idea">💡 Idea</span>'
    if "mvp" in s:    return '<span class="badge-mvp">🔧 MVP</span>'
    if "growth" in s: return '<span class="badge-growth">📈 Growth</span>'
    return f'<span class="badge-mvp">{stage}</span>'


def status_badge(status: str) -> str:
    s = str(status).lower()
    if "active" in s:  return '<span class="badge-active">● Active</span>'
    if "alumni" in s:  return '<span class="badge-alumni">★ Alumni</span>'
    return f'<span class="badge-idea">{status}</span>'


def make_table(headers: list, rows_html: str) -> str:
    th = "".join(
        f'<th style="padding:10px 14px;text-align:{"center" if h in ("Score","Rating","Status") else "left"};'
        f'font-size:11px;color:#475569;font-weight:600;text-transform:uppercase;">{h}</th>'
        for h in headers
    )
    return (
        f'<div style="overflow-x:auto;">'
        f'<table style="width:100%;border-collapse:collapse;border:1px solid #e2e8f0;'
        f'border-radius:10px;overflow:hidden;font-family:inherit;">'
        f'<thead><tr style="background:#f1f5f9;">{th}</tr></thead>'
        f'<tbody>{rows_html}</tbody></table></div>'
    )


# ─────────────────────────────────────────────────────────────────────────────
# RENDER — EXTRACTION PREVIEW
# ─────────────────────────────────────────────────────────────────────────────
def render_extraction_preview(extracted: dict):
    sessions = extracted.get("sessions", [])
    notes    = extracted.get("extraction_notes", "")

    if notes:
        st.markdown(
            f'<div class="ai-insight"><span style="color:#2563eb;font-weight:600;">📖 Extraction Summary: </span>{notes}</div>',
            unsafe_allow_html=True
        )

    if not sessions:
        st.warning("No sessions extracted. Check that your files contain session transcripts.")
        return

    st.markdown(f"**{len(sessions)} session(s) found across files:**")

    type_bg = {"VP Session":"#f5f3ff","Expert Session":"#f0fdf4","Panelist Call":"#fefce8"}
    type_tc = {"VP Session":"#7c3aed","Expert Session":"#059669","Panelist Call":"#d97706"}

    rows_html = ""
    for s in sessions:
        stype   = s.get("type", "")
        bg      = type_bg.get(stype, "#f8fafc")
        tc      = type_tc.get(stype, "#374151")
        probs   = "; ".join(s.get("problems_discussed", []))[:120]
        actions = "; ".join(s.get("action_items", []))[:80]
        rating  = s.get("effectiveness_rating", "—")
        try:
            stars = "★" * int(rating) + "☆" * (5 - int(rating))
        except Exception:
            stars = str(rating)
        rows_html += (
            f'<tr>'
            f'<td style="padding:8px 12px;font-size:12px;color:#374151;">{s.get("session_id","")}</td>'
            f'<td style="padding:8px 12px;">'
            f'<span style="background:{bg};color:{tc};border-radius:6px;padding:2px 8px;font-size:11px;font-weight:600;">{stype}</span>'
            f'</td>'
            f'<td style="padding:8px 12px;font-size:12px;color:#374151;">{s.get("date","")}</td>'
            f'<td style="padding:8px 12px;font-size:12px;color:#374151;">{s.get("participants","")[:50]}</td>'
            f'<td style="padding:8px 12px;font-size:12px;color:#1e293b;">{probs}</td>'
            f'<td style="padding:8px 12px;font-size:12px;color:#374151;">{actions}</td>'
            f'<td style="padding:8px 12px;font-size:13px;color:#d97706;text-align:center;">{stars}</td>'
            f'<td style="padding:8px 12px;font-size:11px;color:#475569;">{s.get("source_file","")}</td>'
            f'</tr>'
        )
    st.markdown(
        make_table(["ID","Type","Date","Participants","Problems","Actions","Rating","Source File"], rows_html),
        unsafe_allow_html=True
    )


# ─────────────────────────────────────────────────────────────────────────────
# RENDER — FULL ANALYSIS
# ─────────────────────────────────────────────────────────────────────────────
def render_analysis(analysis: dict):

    # Deck insights
    di = analysis.get("deck_insights", {})
    if str(di.get("has_deck","")).lower() == "true":
        st.markdown("#### 📋 Pitch Deck Analysis")
        col_d1, col_d2 = st.columns(2)
        with col_d1:
            st.markdown(
                f'<div class="nen-card-accent">'
                f'<div style="font-size:11px;color:#475569;font-weight:600;text-transform:uppercase;margin-bottom:8px;">Problem Statement</div>'
                f'<div style="font-size:14px;color:#1e293b;line-height:1.6;">{di.get("problem_statement","Not available")}</div>'
                f'<div style="height:12px"></div>'
                f'<div style="font-size:11px;color:#475569;font-weight:600;text-transform:uppercase;margin-bottom:8px;">Solution</div>'
                f'<div style="font-size:14px;color:#1e293b;line-height:1.6;">{di.get("solution_summary","Not available")}</div>'
                f'</div>', unsafe_allow_html=True
            )
        with col_d2:
            s_html = "".join(f'<div style="color:#16a34a;font-size:13px;margin-bottom:4px;">✓ {s}</div>' for s in di.get("deck_strengths",[]))
            g_html = "".join(f'<div style="color:#dc2626;font-size:13px;margin-bottom:4px;">✗ {g}</div>' for g in di.get("deck_gaps",[]))
            st.markdown(
                f'<div class="nen-card">'
                f'<div style="font-size:11px;color:#475569;font-weight:600;text-transform:uppercase;margin-bottom:8px;">Strengths</div>{s_html}'
                f'<div style="height:10px"></div>'
                f'<div style="font-size:11px;color:#475569;font-weight:600;text-transform:uppercase;margin-bottom:8px;">Gaps</div>{g_html}'
                f'</div>', unsafe_allow_html=True
            )
        if di.get("deck_vs_reality"):
            st.markdown(f'<div class="ai-insight"><span style="color:#d97706;font-weight:600;">💡 Deck vs Reality: </span>{di["deck_vs_reality"]}</div>', unsafe_allow_html=True)
        st.markdown("---")

    # Health before vs after
    hb = analysis.get("health_before", {})
    ha = analysis.get("health_after",  {})
    sb = hb.get("score", 5)
    sa = ha.get("score", 5)
    st.markdown("#### 📈 Health Score: Before vs After")
    hc1, hc2, hc3 = st.columns([2, 1, 2])
    with hc1:
        st.markdown(
            f'<div class="nen-card" style="text-align:center;">'
            f'<div style="font-size:11px;color:#475569;font-weight:600;text-transform:uppercase;margin-bottom:8px;">Before Program</div>'
            f'<div style="font-size:52px;font-weight:800;color:{health_color(sb)}">{sb}</div>'
            f'<div style="font-size:11px;color:#475569;margin-bottom:8px;">/ 10</div>'
            f'<div style="font-size:12px;color:#374151;line-height:1.5;">{hb.get("rationale","")}</div>'
            f'</div>', unsafe_allow_html=True
        )
    with hc2:
        try:
            delta = int(sa) - int(sb)
        except Exception:
            delta = 0
        arrow = "↑" if delta > 0 else ("↓" if delta < 0 else "→")
        ac = "#16a34a" if delta > 0 else ("#dc2626" if delta < 0 else "#475569")
        st.markdown(
            f'<div style="text-align:center;padding:40px 0;">'
            f'<div style="font-size:32px;color:{ac};font-weight:800;">{arrow}</div>'
            f'<div style="font-size:18px;color:{ac};font-weight:700;">{("+" if delta>0 else "")}{delta}</div>'
            f'<div style="font-size:11px;color:#475569;margin-top:4px;">change</div>'
            f'</div>', unsafe_allow_html=True
        )
    with hc3:
        st.markdown(
            f'<div class="nen-card" style="text-align:center;">'
            f'<div style="font-size:11px;color:#475569;font-weight:600;text-transform:uppercase;margin-bottom:8px;">After All Sessions</div>'
            f'<div style="font-size:52px;font-weight:800;color:{health_color(sa)}">{sa}</div>'
            f'<div style="font-size:11px;color:#475569;margin-bottom:8px;">/ 10</div>'
            f'<div style="font-size:12px;color:#374151;line-height:1.5;">{ha.get("rationale","")}</div>'
            f'</div>', unsafe_allow_html=True
        )
    st.markdown("")

    # Journey narrative
    sp = analysis.get("stage_progression", {})
    mc_map = {"strong":"#16a34a","moderate":"#2563eb","slow":"#d97706","stalled":"#dc2626"}
    m  = sp.get("momentum","moderate")
    mc = mc_map.get(m, "#2563eb")
    st.markdown(
        f'<div class="nen-card-accent">'
        f'<div style="font-size:11px;color:#475569;font-weight:600;text-transform:uppercase;margin-bottom:12px;">Journey Narrative</div>'
        f'<div style="font-size:14px;line-height:1.8;color:#1e3a5f;">{analysis.get("journey_narrative","")}</div>'
        f'<div style="margin-top:14px;display:flex;align-items:center;gap:10px;">'
        f'<span style="font-size:11px;color:#475569;">Momentum:</span>'
        f'<span style="color:{mc};font-weight:600;text-transform:capitalize;">▲ {m}</span>'
        f'<span style="font-size:12px;color:#475569;">— {sp.get("momentum_reason","")}</span>'
        f'</div></div>', unsafe_allow_html=True
    )

    # Key problems table
    st.markdown("#### 🔍 Key Problems Identified")
    problems = analysis.get("key_problems", [])
    if problems:
        status_cfg = {
            "resolved":    ("#f0fdf4","#16a34a","✅"),
            "ongoing":     ("#fefce8","#d97706","⚠️"),
            "unaddressed": ("#fef2f2","#dc2626","🔴"),
        }
        rows_html = ""
        for i, p in enumerate(problems):
            stat = p.get("status","ongoing").lower()
            bg, tc, icon = status_cfg.get(stat, status_cfg["ongoing"])
            row_bg = bg if i % 2 == 0 else "#ffffff"
            rows_html += (
                f'<tr style="background:{row_bg};">'
                f'<td style="padding:10px 14px;font-size:13px;color:#1e293b;line-height:1.5;width:55%;">{p.get("problem","")}</td>'
                f'<td style="padding:10px 14px;font-size:12px;color:#475569;width:25%;">{p.get("source","")}</td>'
                f'<td style="padding:10px 14px;width:20%;text-align:center;">'
                f'<span style="background:{bg};color:{tc};border-radius:20px;padding:3px 10px;font-size:11px;font-weight:600;">{icon} {stat.capitalize()}</span>'
                f'</td></tr>'
            )
        st.markdown(make_table(["Problem","Identified Via","Status"], rows_html), unsafe_allow_html=True)
        st.markdown("")
    else:
        st.info("No problems extracted.")

    # Meeting effectiveness table
    st.markdown("#### 📊 Meeting Effectiveness")
    me_list = analysis.get("meeting_effectiveness", [])
    if me_list:
        type_bg = {"VP Session":"#f5f3ff","Expert Session":"#f0fdf4","Panelist Call":"#fefce8"}
        type_tc = {"VP Session":"#7c3aed","Expert Session":"#059669","Panelist Call":"#d97706"}
        rows_html = ""
        for i, me in enumerate(sorted(me_list, key=lambda x: x.get("effectiveness_score",0), reverse=True)):
            score  = me.get("effectiveness_score", 0)
            mtype  = me.get("type","")
            is_top = me.get("is_most_impactful", False)
            bg     = type_bg.get(mtype,"#f8fafc")
            tc     = type_tc.get(mtype,"#374151")
            row_bg = "#fffbeb" if is_top else ("#f8fafc" if i%2==0 else "#ffffff")
            try:
                stars = "★"*int(score) + "☆"*(5-int(score))
            except Exception:
                stars = str(score)
            trophy = "🏆 " if is_top else ""
            rows_html += (
                f'<tr style="background:{row_bg};">'
                f'<td style="padding:10px 14px;font-size:12px;color:#374151;width:8%;">{me.get("meeting_ref","")}</td>'
                f'<td style="padding:10px 14px;width:18%;">'
                f'<span style="background:{bg};color:{tc};border-radius:6px;padding:3px 8px;font-size:11px;font-weight:700;">{trophy}{mtype}</span>'
                f'</td>'
                f'<td style="padding:10px 14px;font-size:12px;color:#374151;width:12%;">{me.get("date","")}</td>'
                f'<td style="padding:10px 14px;font-size:13px;color:#1e293b;line-height:1.5;">{me.get("impact","")}</td>'
                f'<td style="padding:10px 14px;text-align:center;font-size:15px;color:#d97706;">{stars}</td>'
                f'</tr>'
            )
        st.markdown(make_table(["ID","Type","Date","Impact","Score"], rows_html), unsafe_allow_html=True)
        st.markdown("")

    # Most impactful + priority
    mim = analysis.get("most_impactful_meeting", {})
    if mim:
        st.markdown(
            f'<div class="nen-card-accent">'
            f'<div style="font-size:11px;color:#d97706;text-transform:uppercase;letter-spacing:1px;">🏆 Most Impactful Session</div>'
            f'<div style="font-size:16px;color:#111827;font-weight:700;margin:8px 0;">{mim.get("ref","")} · {mim.get("type","")}</div>'
            f'<div style="color:#374151;font-size:14px;">{mim.get("reason","")}</div>'
            f'</div>', unsafe_allow_html=True
        )
    if analysis.get("next_priority"):
        st.markdown(f'<div class="ai-insight"><span style="color:#2563eb;font-weight:600;">⚡ Top Priority: </span>{analysis["next_priority"]}</div>', unsafe_allow_html=True)
    if analysis.get("biggest_unresolved_problem"):
        st.markdown(f'<div class="ai-insight"><span style="color:#dc2626;font-weight:600;">🔴 Biggest Unresolved Problem: </span>{analysis["biggest_unresolved_problem"]}</div>', unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# RENDER — VENTURE DEEP-DIVE
# ─────────────────────────────────────────────────────────────────────────────
def render_venture_detail(venture_row: pd.Series):
    vid   = str(venture_row.get("Venture_ID",""))
    vname = str(venture_row.get("Venture_Name","Unknown"))

    st.markdown(f"## 🚀 {vname}")
    hub_member = venture_row.get("Hub_Team_Member","—")
    onboard_month  = venture_row.get("Onboarding_Month","—")
    onboard_year   = venture_row.get("Onboarding_Year","—")
    onboard_label  = f"{onboard_month} {onboard_year}" if onboard_month != "—" else "—"
    venture_type   = venture_row.get("Venture_Type","—")
    program_name   = venture_row.get("Program_Name","—")

    # Type badge
    vtype_colors = {"Startup":("#dbeafe","#1e40af"), "SME":("#fef3c7","#92400e")}
    vt_bg, vt_tc = vtype_colors.get(venture_type, ("#f1f5f9","#475569"))
    vtype_badge = f'<span style="background:{vt_bg};color:{vt_tc};border:1px solid {vt_tc};padding:2px 10px;border-radius:20px;font-size:11px;font-weight:600;">{venture_type}</span>'

    # Program badge
    prog_colors = {"Core":("#d1fae5","#065f46"), "Select":("#ede9fe","#5b21b6")}
    pr_bg, pr_tc = prog_colors.get(program_name, ("#f1f5f9","#475569"))
    prog_badge = f'<span style="background:{pr_bg};color:{pr_tc};border:1px solid {pr_tc};padding:2px 10px;border-radius:20px;font-size:11px;font-weight:600;">{program_name}</span>'

    c1, c2, c3, c4, c5, c6, c7 = st.columns(7)
    with c1: st.markdown(f"**Sector** &nbsp; {venture_row.get('Sector','—')}", unsafe_allow_html=True)
    with c2: st.markdown(f"**Founder** &nbsp; {venture_row.get('Founder_Name','—')}", unsafe_allow_html=True)
    with c3: st.markdown(f"**Hub Team** &nbsp; {hub_member}", unsafe_allow_html=True)
    with c4: st.markdown(f"**Onboarded** &nbsp; {onboard_label}", unsafe_allow_html=True)
    with c5: st.markdown(vtype_badge, unsafe_allow_html=True)
    with c6: st.markdown(prog_badge, unsafe_allow_html=True)
    with c7: st.markdown(status_badge(st.session_state.extraction_cache.get(vid,{}).get("current_status","Active")), unsafe_allow_html=True)

    st.markdown(
        f'<div class="nen-card"><em style="color:#374151;">"{venture_row.get("Description","")}"</em></div>',
        unsafe_allow_html=True
    )
    st.markdown(
        f'<span style="color:#475569;font-size:12px;text-transform:uppercase;letter-spacing:1px;">Stage (from files)</span>&nbsp;'
        f'{stage_badge(st.session_state.extraction_cache.get(vid,{}).get("stage_inferred","—"))}'
        f'<span class="journey-arrow"></span>'
        f'<span style="color:#475569;font-size:12px;">— {st.session_state.extraction_cache.get(vid,{}).get("stage_rationale","Extract files to see stage")}</span>'
        f'<span style="color:#475569;font-size:12px;text-transform:uppercase;letter-spacing:1px;">Now</span>&nbsp;'
        f'{stage_badge(str(venture_row.get("Stage_Current","?")))}'
        f'</div>', unsafe_allow_html=True
    )

    # Load Drive files
    folder_path = str(venture_row.get("Folder_Path","")).strip()
    with st.spinner("Loading files from Google Drive…"):
        loaded = load_venture_files(vname, folder_path)

    source_labels = {
        "gdrive": ("☁️ Google Drive","#d1fae5","#065f46"),
        "local":  ("💾 Local","#dbeafe","#1e40af"),
        "none":   ("⚠️ No files","#fef3c7","#92400e"),
    }
    slabel, sbg, stc = source_labels.get(loaded["source"], source_labels["none"])

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Files Found", loaded["total"] if loaded["folder_exists"] else "—")
    c2.metric("Deck Files", len(loaded["deck"]))
    c3.metric("Session Files", loaded["total"] - len(loaded["deck"]))
    with c4:
        st.markdown(
            f'<div style="padding:16px 20px;">'
            f'<div style="font-size:11px;color:#475569;text-transform:uppercase;letter-spacing:1px;margin-bottom:4px;">Source</div>'
            f'<span style="background:{sbg};color:{stc};border-radius:20px;padding:4px 12px;font-size:12px;font-weight:600;">{slabel}</span>'
            f'</div>', unsafe_allow_html=True
        )

    if not loaded["folder_exists"]:
        debug = loaded.get("debug","")
        if "Available:" in debug:
            st.warning(f"⚠️ {debug}")
        elif get_gdrive_service():
            st.warning(f"Folder not found for '{vname}'. {debug}")
        else:
            st.info("💡 Add `gdrive_service_account` to Streamlit secrets to enable Drive file reading.")

    # File browser
    if loaded["files"]:
        type_colors = {"deck":"#d97706","vp":"#7c3aed","expert":"#059669","panelist":"#d97706","other":"#64748b"}
        type_labels = {"deck":"📋 Deck","vp":"🎙 VP","expert":"💡 Expert","panelist":"🏛 Panelist","other":"📄 Other"}
        with st.expander(f"📁 {loaded['total']} file(s) in folder — click to browse", expanded=False):
            chips = "".join(
                f'<span style="background:#f1f5f9;color:{type_colors[ft]};border:1px solid {type_colors[ft]};'
                f'padding:3px 10px;border-radius:20px;font-size:11px;margin-right:6px;font-weight:600;">'
                f'{type_labels[ft]} × {len(loaded[ft])}</span>'
                for ft in ["deck","vp","expert","panelist","other"] if loaded[ft]
            )
            st.markdown(f'<div style="margin-bottom:12px;">{chips}</div>', unsafe_allow_html=True)
            for entry in loaded["files"]:
                tl = type_labels.get(entry["type"],"📄")
                preview = (entry["text"] or "")[:300].replace("\n"," ").strip()
                with st.expander(f"{entry['name']}  ·  {tl}", expanded=False):
                    st.markdown(
                        f'<div style="background:#f8fafc;border:1px solid #e2e8f0;border-radius:8px;'
                        f'padding:12px;font-size:12px;color:#374151;line-height:1.7;white-space:pre-wrap;">'
                        f'{preview}{"…" if len(entry["text"] or "")>300 else ""}</div>',
                        unsafe_allow_html=True
                    )

    st.markdown("---")

    # Step 1: Extract
    st.markdown("### 📄 Step 1 — Extract Sessions from Files")
    if vid not in st.session_state.extraction_cache:
        if loaded["files"]:
            if st.button(f"Extract Sessions for {vname}", key=f"extract_{vid}"):
                with st.spinner("Claude is reading all files and extracting session data…"):
                    try:
                        extracted = ai_extract_sessions(vname, loaded)
                        st.session_state.extraction_cache[vid] = extracted
                        st.rerun()
                    except Exception as e:
                        st.error(f"Extraction failed: {e}")
        else:
            st.info("No files found. Upload session transcripts and deck to the venture's Google Drive folder.")
    else:
        extracted = st.session_state.extraction_cache[vid]
        render_extraction_preview(extracted)

        if st.button("🔄 Re-extract from files", key=f"reextract_{vid}"):
            del st.session_state.extraction_cache[vid]
            st.session_state.analysis_cache.pop(vid, None)
            st.rerun()

        # Step 2: Full analysis
        st.markdown("---")
        st.markdown("### 🤖 Step 2 — Full AI Analysis")
        n_sessions = len(extracted.get("sessions",[]))
        has_deck   = len(loaded["deck"]) > 0

        if vid not in st.session_state.analysis_cache:
            label = f"Generate Full Analysis  ({n_sessions} sessions · {'deck ✓' if has_deck else 'no deck'})"
            if st.button(label, key=f"analyze_{vid}"):
                with st.spinner("Generating venture intelligence report…"):
                    try:
                        analysis = ai_analyze_venture(venture_row, loaded, extracted)
                        st.session_state.analysis_cache[vid] = analysis
                        st.rerun()
                    except Exception as e:
                        st.error(f"Analysis failed: {e}")
        else:
            if st.button("🔄 Re-analyse", key=f"reanalyze_{vid}"):
                st.session_state.analysis_cache.pop(vid, None)
                st.rerun()
            else:
                render_analysis(st.session_state.analysis_cache[vid])


# ─────────────────────────────────────────────────────────────────────────────
# RENDER — PORTFOLIO
# ─────────────────────────────────────────────────────────────────────────────
def render_portfolio_view(ventures_df: pd.DataFrame):
    st.markdown("## 📊 Portfolio Overview")
    total    = len(ventures_df)
    active   = sum(1 for vid in st.session_state.extraction_cache if st.session_state.extraction_cache[vid].get("current_status","") == "Active")
    alumni   = sum(1 for vid in st.session_state.extraction_cache if st.session_state.extraction_cache[vid].get("current_status","") == "Alumni")
    alumni   = len(ventures_df[ventures_df["Status"]=="Alumni"]) if "Status" in ventures_df.columns else 0
    analysed = len(st.session_state.analysis_cache)

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Total Ventures", total)
    c2.metric("Active", active)
    c3.metric("Alumni", alumni)
    c4.metric("Analysed", f"{analysed}/{total}")
    st.markdown("---")

    st.markdown("### All Ventures")
    for _, row in ventures_df.iterrows():
        vid     = str(row.get("Venture_ID",""))
        cached  = st.session_state.analysis_cache.get(vid, {})
        hb      = cached.get("health_before",{}).get("score","—")
        ha      = cached.get("health_after", {}).get("score","—")
        n_sess  = len(st.session_state.extraction_cache.get(vid,{}).get("sessions",[]))
        try:
            d = int(ha) - int(hb)
            ac = "#16a34a" if d>0 else ("#dc2626" if d<0 else "#475569")
            delta_html = f'<span style="color:{ac};font-weight:700;margin-left:4px;">{("+" if d>0 else "")}{d}</span>'
        except Exception:
            delta_html = ""

        st.markdown(
            f'<div class="nen-card">'
            f'<div style="display:flex;justify-content:space-between;align-items:flex-start;flex-wrap:wrap;gap:10px;">'
            f'<div><div style="font-weight:700;font-size:16px;color:#111827;">{row.get("Venture_Name","")}</div>'
            f'<div style="color:#475569;font-size:12px;margin-top:2px;">{row.get("Founder_Name","")} · {row.get("Sector","")} · <span style="color:#2563eb;">👤 {row.get("Hub_Team_Member","—")}</span> · 📅 {row.get("Onboarding_Month","—")} {row.get("Onboarding_Year","")} · {row.get("Venture_Type","—")} · {row.get("Program_Name","—")}</div></div>'
            f'<div style="display:flex;align-items:center;gap:10px;flex-wrap:wrap;">'
            f'{stage_badge(st.session_state.extraction_cache.get(vid,{}).get("stage_inferred","—"))}'
            f'&nbsp;&nbsp;{status_badge(st.session_state.extraction_cache.get(vid,{}).get("current_status","Active"))}'
            f'</div></div>'
            f'<div style="display:flex;gap:20px;margin-top:12px;flex-wrap:wrap;">'
            f'<span style="font-size:12px;color:#475569;">Sessions extracted: <b>{n_sess}</b></span>'
            f'<span style="font-size:12px;color:#475569;">Health before: <b>{hb}</b></span>'
            f'<span style="font-size:12px;color:#475569;">After: <b>{ha}</b>{delta_html}</span>'
            f'</div></div>',
            unsafe_allow_html=True
        )

    st.markdown("---")
    st.markdown("### 🤖 Portfolio AI Analysis")
    if "portfolio_analysis" not in st.session_state:
        if st.button("Generate Portfolio Intelligence Report"):
            with st.spinner("Analyzing portfolio…"):
                try:
                    st.session_state.portfolio_analysis = ai_portfolio_analysis(ventures_df)
                    st.rerun()
                except Exception as e:
                    st.error(f"Portfolio analysis failed: {e}")
    else:
        pa = st.session_state.portfolio_analysis
        ph = pa.get("portfolio_health",{})
        hs = ph.get("score", 5)
        col1, col2 = st.columns([1,3])
        with col1:
            st.markdown(
                f'<div class="nen-card" style="text-align:center;">'
                f'<div style="font-size:11px;color:#475569;font-weight:600;text-transform:uppercase;margin-bottom:8px;">Portfolio Health</div>'
                f'<div style="font-size:52px;font-weight:800;color:{health_color(hs)}">{hs}</div>'
                f'<div style="font-size:11px;color:#475569;">/ 10</div>'
                f'</div>', unsafe_allow_html=True
            )
        with col2:
            st.markdown(
                f'<div class="nen-card-accent">'
                f'<div style="font-size:11px;color:#475569;font-weight:600;text-transform:uppercase;margin-bottom:10px;">Cohort Story</div>'
                f'<div style="font-size:14px;line-height:1.8;color:#1e3a5f;">{pa.get("cohort_narrative","")}</div>'
                f'</div>', unsafe_allow_html=True
            )
        col1, col2 = st.columns(2)
        with col1:
            st.markdown("#### ⚠️ Needs Attention")
            for item in pa.get("needs_attention",[]):
                st.markdown(
                    f'<div style="background:#fef2f2;border:1px solid #fca5a5;border-radius:8px;padding:10px 14px;margin-bottom:8px;">'
                    f'<div style="color:#dc2626;font-weight:600;">{item.get("name","")}</div>'
                    f'<div style="color:#374151;font-size:12px;margin-top:4px;">{item.get("reason","")}</div>'
                    f'</div>', unsafe_allow_html=True
                )
        with col2:
            st.markdown("#### 🌟 Best Performing")
            for item in pa.get("best_performing",[]):
                st.markdown(
                    f'<div style="background:#f0fdf4;border:1px solid #86efac;border-radius:8px;padding:10px 14px;margin-bottom:8px;">'
                    f'<div style="color:#16a34a;font-weight:600;">{item.get("name","")}</div>'
                    f'<div style="color:#374151;font-size:12px;margin-top:4px;">{item.get("reason","")}</div>'
                    f'</div>', unsafe_allow_html=True
                )
        mte = pa.get("meeting_type_effectiveness",{})
        if mte:
            st.markdown("#### 🏆 Most Impactful Meeting Type")
            st.markdown(
                f'<div class="nen-card-accent"><div style="display:flex;gap:40px;flex-wrap:wrap;">'
                f'<div><div style="font-size:11px;color:#475569;font-weight:600;text-transform:uppercase;">Most Effective</div>'
                f'<div style="font-size:20px;font-weight:700;color:#16a34a;margin:6px 0;">{mte.get("most_effective_type","")}</div>'
                f'<div style="font-size:13px;color:#374151;">{mte.get("reason","")}</div></div>'
                f'<div><div style="font-size:11px;color:#475569;font-weight:600;text-transform:uppercase;">Least Effective</div>'
                f'<div style="font-size:20px;font-weight:700;color:#dc2626;margin:6px 0;">{mte.get("least_effective_type","")}</div>'
                f'<div style="font-size:13px;color:#374151;">{mte.get("least_effective_reason","")}</div></div>'
                f'</div></div>', unsafe_allow_html=True
            )
        st.markdown("#### 💡 Portfolio Recommendations")
        for i, rec in enumerate(pa.get("portfolio_recommendations",[]), 1):
            st.markdown(f'<div class="ai-insight"><span style="color:#2563eb;font-weight:700;">{i}.</span> {rec}</div>', unsafe_allow_html=True)
        if st.button("🔄 Regenerate Portfolio Report"):
            del st.session_state.portfolio_analysis
            st.rerun()


# ─────────────────────────────────────────────────────────────────────────────
# CHAT
# ─────────────────────────────────────────────────────────────────────────────
def render_chat(ventures_df: pd.DataFrame, selected_name: str = None, view_mode: str = None):
    st.markdown("---")
    st.markdown("### 💬 Ask Anything")

    for msg in st.session_state.chat_messages:
        with st.chat_message(msg["role"]):
            st.markdown(msg["content"])

    if not st.session_state.chat_messages:
        if view_mode == "Venture Deep-Dive" and selected_name:
            hint = f"*Try: What are {selected_name}'s biggest problems? · Which session helped most? · Summarise the journey*"
        else:
            hint = "*Try: Which ventures need most attention? · Which meeting type was most effective? · Compare all ventures*"
        st.markdown(f'<div style="color:#64748b;font-size:13px;padding:8px 0;">{hint}</div>', unsafe_allow_html=True)

    user_input = st.chat_input("Ask about a venture, sessions, or the portfolio…")
    if user_input:
        with st.chat_message("user"):
            st.markdown(user_input)
        st.session_state.chat_messages.append({"role":"user","content":user_input})

        with st.chat_message("assistant"):
            with st.spinner("Thinking…"):
                try:
                    # Build context
                    ctx = "VENTURES:\n"
                    for _, r in ventures_df.iterrows():
                        ctx += f"- {r.get('Venture_Name','')} | {r.get('Sector','')} | {r.get('Stage_Start','')}→{r.get('Stage_Current','')} | {r.get('Status','')}\n"
                    for vid, ext in st.session_state.extraction_cache.items():
                        vrow = ventures_df[ventures_df.get("Venture_ID","") == vid] if "Venture_ID" in ventures_df.columns else pd.DataFrame()
                        vname = vrow.iloc[0].get("Venture_Name", vid) if not vrow.empty else vid
                        for s in ext.get("sessions",[]):
                            ctx += f"  [{vname}] {s.get('type','')} {s.get('date','')} — {s.get('notes','')[:80]}\n"
                    for vid, an in st.session_state.analysis_cache.items():
                        vrow = ventures_df[ventures_df.get("Venture_ID","") == vid] if "Venture_ID" in ventures_df.columns else pd.DataFrame()
                        vname = vrow.iloc[0].get("Venture_Name", vid) if not vrow.empty else vid
                        hb = an.get("health_before",{}).get("score","?")
                        ha = an.get("health_after", {}).get("score","?")
                        ctx += f"  [{vname}] Health {hb}→{ha} | {an.get('next_priority','')}\n"

                    history = [{"role": m["role"], "content": m["content"]}
                               for m in st.session_state.chat_messages[-8:]]
                    msgs = history[:-1] + [{"role":"user","content":f"DATA:\n{ctx}\n\nQUESTION: {user_input}"}]
                    resp = client.messages.create(
                        model="claude-sonnet-4-20250514",
                        max_tokens=1000,
                        system="You are an analyst assistant for NEN Resources Network. Answer questions about ventures and sessions concisely using the data provided.",
                        messages=msgs
                    )
                    reply = resp.content[0].text
                    st.markdown(reply)
                    st.session_state.chat_messages.append({"role":"assistant","content":reply})
                except Exception as e:
                    st.error(f"Chat error: {e}")


# ─────────────────────────────────────────────────────────────────────────────
# HEADER
# ─────────────────────────────────────────────────────────────────────────────
hcol1, hcol2, _ = st.columns([1, 2, 1])
with hcol1:
    if os.path.exists("DP_BG1.png"):
        st.image("DP_BG1.png", width=150)
    else:
        st.markdown('<div style="font-weight:800;font-size:22px;color:#111827;padding:10px 0;">NEN</div>', unsafe_allow_html=True)
with hcol2:
    st.markdown("<h2 style='text-align:center;font-weight:800;color:#111827;margin:0;padding:8px 0;'>🚀 Resources Network — Venture Intelligence</h2>", unsafe_allow_html=True)
st.markdown("---")


# ─────────────────────────────────────────────────────────────────────────────
# SIDEBAR
# ─────────────────────────────────────────────────────────────────────────────
with st.sidebar:
    if os.path.exists("DP_BG1.png"):
        st.image("DP_BG1.png", width=130)
    st.markdown('<div style="font-size:11px;color:#64748b;text-transform:uppercase;letter-spacing:1.5px;margin:8px 0 16px;">Portfolio Analyzer</div>', unsafe_allow_html=True)
    st.markdown("---")

    uploaded = st.file_uploader(
        "📊 Upload Ventures Excel",
        type=["xlsx","xls"],
        help="Excel with Ventures sheet: Venture_ID, Venture_Name, Founder_Name, Founder_Email, Sector, Hub_Team_Member. All session data comes from Google Drive."
    )

    if uploaded:
        ventures_df = load_excel(uploaded.read())
        if not ventures_df.empty:
            st.success(f"✓ {len(ventures_df)} ventures loaded")
        else:
            st.error("Could not read ventures. Check sheet name is 'Ventures'.")

        st.markdown("---")

        # ── Hub Team filter ───────────────────────────────────────────────
        filtered_df = ventures_df.copy()
        selected_hub = None
        if "Hub_Team_Member" in ventures_df.columns:
            hub_members = sorted(ventures_df["Hub_Team_Member"].dropna().unique().tolist())
            hub_options = ["All Ventures"] + hub_members
            selected_hub = st.selectbox(
                "👤 Hub Team Member",
                hub_options,
                help="Select your name to see only your ventures"
            )
            if selected_hub and selected_hub != "All Ventures":
                filtered_df = ventures_df[ventures_df["Hub_Team_Member"] == selected_hub].copy()

        # ── Venture Type filter ───────────────────────────────────────────
        if "Venture_Type" in ventures_df.columns:
            vtype_opts = ["All Types"] + sorted(ventures_df["Venture_Type"].dropna().unique().tolist())
            selected_vtype = st.selectbox("🏢 Venture Type", vtype_opts)
            if selected_vtype != "All Types":
                filtered_df = filtered_df[filtered_df["Venture_Type"] == selected_vtype].copy()

        # ── Program filter ────────────────────────────────────────────────
        if "Program_Name" in ventures_df.columns:
            prog_opts = ["All Programs"] + sorted(ventures_df["Program_Name"].dropna().unique().tolist())
            selected_prog = st.selectbox("📋 Program", prog_opts)
            if selected_prog != "All Programs":
                filtered_df = filtered_df[filtered_df["Program_Name"] == selected_prog].copy()

        # Show filter summary
        if len(filtered_df) < len(ventures_df):
            st.info(f"Showing {len(filtered_df)} of {len(ventures_df)} ventures")

        st.markdown("---")
        view_mode = st.selectbox("View Mode", ["Portfolio Overview", "Venture Deep-Dive"])
        selected_name = None
        if view_mode == "Venture Deep-Dive":
            if "Venture_Name" in filtered_df.columns and not filtered_df.empty:
                selected_name = st.selectbox("Select Venture", filtered_df["Venture_Name"].tolist())
            elif filtered_df.empty:
                st.warning(f"No ventures assigned to {selected_hub}.")
            else:
                st.warning("No Venture_Name column found.")

        st.markdown("---")
        if st.sidebar.button("🗑️ Clear Chat & Cache"):
            st.session_state.chat_messages = []
            st.session_state.analysis_cache = {}
            st.session_state.extraction_cache = {}
            st.session_state.pop("portfolio_analysis", None)
            st.rerun()
    else:
        ventures_df = None
        filtered_df = None
        view_mode = None
        selected_name = None
        selected_hub = None

    st.markdown("---")
    st.markdown('<div style="font-size:10px;color:#94a3b8;text-align:center;">Powered by Claude AI · NEN Resources Network</div>', unsafe_allow_html=True)


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────
if ventures_df is None or ventures_df.empty or filtered_df is None:
    st.markdown(
        '<div style="text-align:center;padding:40px 20px 20px;">'
        '<div style="font-weight:800;font-size:38px;color:#111827;line-height:1.1;margin-bottom:16px;">'
        'Venture Intelligence<br><span style="color:#2563eb;">Portfolio Analyzer</span></div>'
        '<div style="font-size:16px;color:#64748b;max-width:520px;margin:0 auto 32px;">'
        'Upload your Ventures Excel to get started. All session data is read directly from Google Drive.'
        '</div></div>', unsafe_allow_html=True
    )
    col1, col2, col3 = st.columns(3)
    for col, icon, title, desc in [
        (col1, "☁️", "Drive-First", "Transcripts and notes read directly from your Google Drive venture folders — no manual data entry"),
        (col2, "🤖", "Two-Step AI", "Step 1: Extract sessions from files. Step 2: Full analysis with before/after health scores and meeting effectiveness"),
        (col3, "💬", "Chat Interface", "Ask anything about any venture or the whole portfolio in plain language"),
    ]:
        with col:
            st.markdown(
                f'<div class="nen-card" style="text-align:center;padding:28px 20px;">'
                f'<div style="font-size:28px;margin-bottom:10px;">{icon}</div>'
                f'<div style="font-weight:700;color:#111827;margin-bottom:6px;">{title}</div>'
                f'<div style="font-size:13px;color:#475569;">{desc}</div>'
                f'</div>', unsafe_allow_html=True
            )

    st.markdown("---")
    st.markdown("### How It Works")
    st.markdown(
        '<div class="nen-card">'
        '<div style="font-weight:700;color:#2563eb;margin-bottom:8px;">Excel — basics only (one sheet: Ventures)</div>'
        '<div style="font-size:13px;color:#374151;line-height:2;">Venture_ID &nbsp;·&nbsp; '
        '<b>Venture_Name</b> <span style="color:#d97706;">(must match Google Drive folder name exactly)</span> &nbsp;·&nbsp; '
        'Founder_Name &nbsp;·&nbsp; Founder_Email &nbsp;·&nbsp; Sector &nbsp;·&nbsp; Stage_Start &nbsp;·&nbsp; Stage_Current &nbsp;·&nbsp; Cohort &nbsp;·&nbsp; Status &nbsp;·&nbsp; Description</div>'
        '<div style="margin-top:12px;font-size:13px;color:#16a34a;font-weight:600;">✓ No session sheets — all session data comes from files in Google Drive</div>'
        '</div>',
        unsafe_allow_html=True
    )
    st.markdown(
        '<div class="ai-insight">'
        '<span style="color:#2563eb;font-weight:600;">📁 Google Drive folder naming: </span>'
        'Venture_Docs/ → Qarmatek/ → files here &nbsp;·&nbsp; '
        'Keywords: <b>vp / meeting / call</b> → VP session &nbsp;·&nbsp; '
        '<b>expert / session / mentor</b> → Expert session &nbsp;·&nbsp; '
        '<b>panelist / panel / investor</b> → Panelist call &nbsp;·&nbsp; '
        '<b>deck / pitch</b> → Pitch deck'
        '</div>',
        unsafe_allow_html=True
    )

else:
    # Show hub filter context banner if filtered
    if selected_hub and selected_hub != "All Ventures":
        st.markdown(
            f'<div style="background:#eff6ff;border:1px solid #bfdbfe;border-radius:10px;'
            f'padding:10px 18px;margin-bottom:16px;display:flex;align-items:center;gap:10px;">'
            f'<span style="font-size:18px;">👤</span>'
            f'<span style="color:#1e40af;font-weight:600;">{selected_hub}</span>'
            f'<span style="color:#475569;font-size:13px;">— showing {len(filtered_df)} assigned venture(s)</span>'
            f'</div>',
            unsafe_allow_html=True
        )

    if view_mode == "Portfolio Overview":
        render_portfolio_view(filtered_df)
        render_chat(filtered_df, None, "Portfolio Overview")

    elif view_mode == "Venture Deep-Dive" and selected_name:
        venture_row = filtered_df[filtered_df["Venture_Name"] == selected_name].iloc[0]
        render_venture_detail(venture_row)
        render_chat(filtered_df, selected_name, "Venture Deep-Dive")
